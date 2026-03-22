import io
import os
import re
import asyncio
import hashlib
import logging
import tempfile
from typing import Any, AsyncIterator, Iterator, Optional

import google.generativeai as genai
import pypdf
from google.generativeai.types import HarmBlockThreshold, HarmCategory
from PIL import Image
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

# Gemini Vision（OCR）用。学術・事務文書で過度にブロックされない設定
VISION_SAFETY_SETTINGS = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

GEMINI_VISION_MODEL = "models/gemini-2.5-flash"

_genai_vision_configured = False


def _ensure_genai_vision_configured() -> None:
    """OCR 用に Gemini を環境変数の API キーで初期化（冪等）。"""
    global _genai_vision_configured
    if _genai_vision_configured:
        return
    genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
    _genai_vision_configured = True


def _gemini_response_text(response: Any) -> str:
    """generate_content の応答からテキストを取り出す（ブロック時のフォールバック含む）。"""
    try:
        t = getattr(response, "text", None)
        if t and str(t).strip():
            return str(t).strip()
    except Exception:
        pass
    try:
        cands = getattr(response, "candidates", None) or []
        if not cands:
            return ""
        parts = getattr(cands[0].content, "parts", None) or []
        return "".join(getattr(p, "text", "") or "" for p in parts).strip()
    except Exception:
        return ""


OCR_VISION_PROMPT_TEMPLATE = """この画像は公的・業務用の文書資料の1ページです。次を厳守して文字起こししてください。

【出力形式】
- 必ず Markdown で構造化すること（見出しは # / ## / ###、箇条書きは - または番号付きリスト）。
- 表は Markdown テーブル（| 列1 | 列2 |）で再現し、列・行の対応を維持すること。結合セルは可能な範囲で注記で示すこと。
- 脚注・注釈・欄外注記・図表のキャプション・※印の説明は省略せず、本文と区別できる形で保持すること。
- 読み取れない箇所のみ [判読不能] とし、それ以外は推測で補完しないこと。
- 前置き・後書き・「この画像は〜」などのメタ説明は出力しない。抽出した本文のみを返すこと。

【文脈】元PDFファイル名: {source_pdf_name} / ページ番号: {page_num}
"""


async def _extract_text_with_gemini_vision(
    jpeg_bytes: bytes,
    page_num: int,
    source_pdf_name: str,
    max_retries: int = 4,
) -> str:
    """
    Gemini 2.5 Flash による画像OCR（Markdown構造化指示）。
    429 / Quota 時は asyncio.sleep しながらリトライする。
    """
    _ensure_genai_vision_configured()
    if not os.getenv("GEMINI_API_KEY"):
        logging.error("GEMINI_API_KEY が未設定のため PDF OCR を実行できません")
        return ""

    prompt = OCR_VISION_PROMPT_TEMPLATE.format(
        source_pdf_name=source_pdf_name,
        page_num=page_num,
    )
    image = Image.open(io.BytesIO(jpeg_bytes))
    model = genai.GenerativeModel(GEMINI_VISION_MODEL)

    last_err: Optional[Exception] = None
    for attempt in range(max_retries + 1):
        try:
            response = await model.generate_content_async(
                [prompt, image],
                safety_settings=VISION_SAFETY_SETTINGS,
            )
            text = _gemini_response_text(response)
            if text:
                return text
            logging.warning(
                "Gemini Vision が空応答 (page=%s, attempt=%s)", page_num, attempt
            )
            if attempt < max_retries:
                await asyncio.sleep(2)
                continue
            return ""
        except Exception as e:
            last_err = e
            err_s = str(e).lower()
            is_quota = "429" in str(e) or "quota" in err_s or "resource_exhausted" in err_s
            if is_quota and attempt < max_retries:
                wait = 25 + attempt * 15
                logging.warning(
                    "Gemini Vision 429/Quota (page=%s)。%s秒待機してリトライ (%s/%s)",
                    page_num,
                    wait,
                    attempt + 1,
                    max_retries,
                )
                await asyncio.sleep(wait)
                continue
            logging.error("Gemini Vision OCR エラー (page=%s): %s", page_num, e)
            if attempt < max_retries:
                await asyncio.sleep(5)
                continue
            break
    if last_err:
        logging.error("Gemini Vision OCR 最終失敗 (page=%s): %s", page_num, last_err)
    return ""


class SimpleDocumentProcessor:
    """
    Render (メモリ512MB) 環境向け：メモリ効率重視の親子チャンキング対応プロセッサー
    特徴:
    1. DOCXなどの構造化ファイルをMarkdown形式に変換して抽出。
    2. 全データをリストで保持せず、Iterator (yield) で1件ずつ返す。
    3. 親子チャンキングのメタデータに「親の文脈」を埋め込む。
    """
    def __init__(self, parent_chunk_size=1500, child_chunk_size=400):
        # Markdown構造を意識したセパレーター設定
        # 見出し(###)や改行(\n\n)を優先して区切ることで、文脈の分断を防ぐ
        separators = [
            "\n#{1,6} ", # Markdownの見出し (# H1, ## H2...)
            "```\n",     # コードブロック
            "\n\n",      # 段落
            "\n",        # 行
            "。", "、", " ", ""
        ]

        # 親チャンク用スプリッター（文脈保持用）
        self.parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=parent_chunk_size,
            chunk_overlap=0,
            separators=separators
        )
        # 子チャンク用スプリッター（検索用ベクトル生成用）
        self.child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=child_chunk_size,
            chunk_overlap=50,
            separators=separators
        )
        logging.info(f"SimpleDocumentProcessor Initialized (Markdown Mode): Parent={parent_chunk_size}, Child={child_chunk_size}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        """DOCXのスタイル情報を読み取り、Markdown形式のテキストに変換する"""
        try:
            doc = DocxDocument(io.BytesIO(content))
            full_text = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name.lower()
                # スタイルをMarkdownにマッピング
                if 'heading 1' in style_name:
                    full_text.append(f"# {text}")
                elif 'heading 2' in style_name:
                    full_text.append(f"## {text}")
                elif 'heading 3' in style_name:
                    full_text.append(f"### {text}")
                elif 'list' in style_name or 'bullet' in style_name:
                    full_text.append(f"- {text}")
                else:
                    # 通常の段落
                    full_text.append(text)
            # 段落ごとに改行を2つ入れて結合
            return "\n\n".join(full_text)
        except Exception as e:
            logging.error(f"DOCX extraction error: {e}")
            return ""

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """PDFからテキスト抽出（PyPDF2使用、メモリ重視）"""
        text = ""
        try:
            # pypdfを使用するように修正
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text
        except Exception as e:
            logging.error(f"PDF extraction error: {e}")
            return ""

    def _extract_text(self, filename: str, content: bytes) -> str:
        """ファイル拡張子に応じて適切な抽出メソッドを呼び出す"""
        text = ""
        try:
            if filename.endswith(".docx"):
                text = self._extract_text_from_docx(content)
                logging.info(f".docx parsed to Markdown: {len(text)} chars")
            elif filename.endswith(".pdf"):
                text = self._extract_text_from_pdf(content)
                # PDFはMarkdown変換が難しいため、基本的なクリーニングのみ
                text = self._clean_text(text)
                logging.info(f".pdf extracted: {len(text)} chars")
            elif filename.endswith(".txt") or filename.endswith(".md"):
                text = content.decode('utf-8')
                logging.info(f".txt/.md loaded: {len(text)} chars")
            else:
                logging.warning(f"Unsupported file type: {filename}")
            return text
        except Exception as e:
            logging.error(f"General text extraction error ({filename}): {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        """汎用的なテキストクリーニング"""
        # 1. 連続する空白を1つに
        text = re.sub(r'[ \t\u3000]+', ' ', text)
        # 2. 3つ以上の連続改行を2つ（段落区切り）に統一
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def process_and_chunk(self, filename: str, content: bytes, category: str, collection_name: str) -> Iterator[LangChainDocument]:
        """
        ListではなくIteratorを返す（メモリ節約）。
        DOCX等の場合はMarkdown形式でテキストが渡ってくるため、構造を維持したままチャンキングされる。
        """
        # 1. テキスト抽出 (DOCXならMarkdown化されている)
        full_text = self._extract_text(filename, content)
        if not full_text:
            return

        # 2. 親チャンク（大きな塊）を生成
        # Markdownヘッダーなどを考慮したスプリッター設定により、章の途中で切れにくくなる
        parent_chunks = self.parent_splitter.split_text(full_text)
        full_text = None # メモリ解放

        logging.info(f"{filename}: {len(parent_chunks)} parent chunks generated.")

        # 3. 各親チャンクをループ処理して子チャンクを作る
        for parent_idx, parent_text in enumerate(parent_chunks):
            # 親の中から子チャンクを生成
            child_chunks = self.child_splitter.split_text(parent_text)
            # document_processor.py の process_and_chunk メソッド内

            for child_text in child_chunks:
                metadata = {
                    "source": filename,
                    "collection_name": collection_name,
                    "category": category,
                    "element_type": "ParentChild",
                    "parent_index": parent_idx,
                    # ★修正: キー名を "parent_context" から "parent_content" に統一
                    # これにより、search.py や documents.py での扱いが統一されます
                    "parent_content": parent_text
                }
                yield LangChainDocument(page_content=child_text, metadata=metadata)
        logging.info(f"{filename}: Processing complete (Stream finished).")


# ---------------------------------------------------------------------------
# 高精度PDFモード: 1ページずつ画像化 → Storage へ直接アップロード → Gemini Vision OCR → チャンキング
# convert_from_path に poppler_path は渡さない（システム PATH の poppler を使用）
# ---------------------------------------------------------------------------


def _pdf_page_count(pdf_path: str) -> int:
    """pdf2image: 総ページ数（1ページ以上）。"""
    from pdf2image import pdfinfo_from_path

    info = pdfinfo_from_path(pdf_path)
    raw = info.get("Pages", 1)
    try:
        n = int(raw)
    except (TypeError, ValueError):
        n = 1
    return max(1, n)


def _render_single_page_jpeg(pdf_path: str, page_num: int, dpi: int = 150) -> bytes:
    """1ページのみレンダリングし JPEG バイナリにする（全ページ一括禁止）。"""
    from pdf2image import convert_from_path

    images = convert_from_path(
        pdf_path,
        first_page=page_num,
        last_page=page_num,
        dpi=dpi,
    )
    try:
        img = images[0]
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        return buf.getvalue()
    finally:
        images.clear()
        del images


def _upload_jpeg_to_supabase(
    supabase_client: Any,
    bucket: str,
    storage_object_path: str,
    jpeg_bytes: bytes,
) -> None:
    """Supabase Storage に JPEG を直接アップロード（ローカルファイルに保存しない）。"""
    supabase_client.storage.from_(bucket).upload(
        storage_object_path,
        jpeg_bytes,
        {"content-type": "image/jpeg"},
    )


async def iter_pdf_ocr_document_chunks(
    pdf_bytes: bytes,
    original_filename: str,
    category: str,
    collection_name: str,
    supabase_client: Any,
    processor: SimpleDocumentProcessor,
    bucket: str = "images",
    dpi: int = 150,
) -> AsyncIterator[LangChainDocument]:
    """
    PDF を一時ファイルに保存し、ページ単位で画像化→Storage アップロード→Gemini Vision OCR→SimpleDocumentProcessor でチャンキング。
    各チャンクの metadata に element_type=image_ocr, image_path（ハッシュ付きjpgファイル名）, source を付与する。
    親子チャンキングの親文脈キーは process_and_chunk 内の parent_content に統一。
    """
    safe_base = re.sub(r"[^\w\.\-]", "_", os.path.basename(original_filename))[:120] or "document"
    fd, path = tempfile.mkstemp(suffix=".pdf")
    try:
        os.write(fd, pdf_bytes)
        os.close(fd)
        n_pages = await asyncio.to_thread(_pdf_page_count, path)
        for page_num in range(1, n_pages + 1):
            jpeg_bytes = await asyncio.to_thread(
                _render_single_page_jpeg, path, page_num, dpi
            )
            digest = hashlib.sha256(jpeg_bytes).hexdigest()[:24]
            # Storage に保存するオブジェクト名（ハッシュ付き・一意）— metadata.image_path に同じ値を格納
            image_filename = f"{safe_base}_p{page_num}_{digest}.jpg"
            storage_object_path = f"pdf_ocr/{image_filename}"
            await asyncio.to_thread(
                _upload_jpeg_to_supabase,
                supabase_client,
                bucket,
                storage_object_path,
                jpeg_bytes,
            )
            ocr_text = await _extract_text_with_gemini_vision(
                jpeg_bytes, page_num, original_filename
            )
            if not ocr_text.strip():
                ocr_text = (
                    f"[ページ{page_num}] OCR結果が空でした。"
                    f"画像は Storage に保存済み: {image_filename}"
                )
            synthetic_filename = f"{safe_base}_p{page_num}.txt"
            for doc in processor.process_and_chunk(
                synthetic_filename,
                ocr_text.encode("utf-8"),
                category,
                collection_name,
            ):
                doc.metadata["element_type"] = "image_ocr"
                doc.metadata["image_path"] = image_filename
                doc.metadata["source"] = original_filename
                # 親子チャンキングの親テキストキーは parent_content に統一（上書きしない）
                yield doc
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


# グローバルインスタンス
simple_processor: Optional[SimpleDocumentProcessor] = None