# --------------------------------------------------------------------------
# 1. ライブラリのインポート
# --------------------------------------------------------------------------
from typing import List, Dict,Any
from collections import defaultdict
import logging
import uvicorn
import traceback
import csv

import uuid
import io

import asyncio
import re

from contextlib import asynccontextmanager

# --- サードパーティライブラリ ---
import google.generativeai as genai
from google.generativeai.types import GenerationConfig

from google.generativeai.types import GenerationConfig
# ★★★ 以下の2行を追加 ★★★
from google.generativeai.types import HarmCategory, HarmBlockThreshold

from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form, WebSocket, WebSocketDisconnect, Depends, Query
from fastapi.responses import HTMLResponse, StreamingResponse, RedirectResponse, FileResponse, Response
from pydantic import BaseModel
# ★ (修正) 低品質アップロード用に PyPDF2 と python-docx を復活
import PyPDF2
from docx import Document as DocxDocument
# (BeautifulSoup は /scrape を削除したため不要)
from starlette.middleware.sessions import SessionMiddleware

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument
from core.config import GEMINI_API_KEY, APP_SECRET_KEY, SUPABASE_URL, SUPABASE_KEY

# --- 初期設定 ---
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(asctime)s - %(message)s')

# --------------------------------------------------------------------------
# 2. 環境変数と基本設定
# --------------------------------------------------------------------------
genai.configure(api_key=GEMINI_API_KEY)
# キーワードマッピング

# データベースから読み込むフォールバック情報を格納するグローバル変数

# --------------------------------------------------------------------------
# 3. 内部コンポーネントの定義
# --------------------------------------------------------------------------

# (split_text 関数 削除済み)

def format_urls_as_links(text: str) -> str:
    """URLをHTMLリンクに変換"""
    url_pattern = r'(?<!\]\()(https?://[^\s\[\]()<>]+)(?!\))'
    text = re.sub(url_pattern, r'<a href="\1" target="_blank">\1</a>', text)
    md_link_pattern = r'\[([^\]]+)\]\((https?://[^\s\)]+)\)'
    text = re.sub(md_link_pattern, r'<a href="\2" target="_blank">\1</a>', text)
    return text

async def safe_generate_content(model, prompt, stream=False, max_retries=3):
    """レート制限を考慮した安全なコンテンツ生成"""
    for attempt in range(max_retries):
        try:
            if stream:
                return await model.generate_content_async(
                    prompt,
                    stream=True,
                    generation_config=GenerationConfig(
                        max_output_tokens=1024,
                        temperature=0.1
                    )
                )
            else:
                return await model.generate_content_async(
                    prompt,
                    generation_config=GenerationConfig(
                        max_output_tokens=512,
                        temperature=0.3
                    )
                )
        except Exception as e:
            error_str = str(e)
            if "429" in error_str or "quota" in error_str.lower():
                if attempt < max_retries - 1:
                    wait_time = 15
                    if "retry in" in error_str:
                        try:
                            match = re.search(r'retry in (\d+(?:\.\d+)?)s', error_str)
                            if match:
                                wait_time = float(match.group(1)) + 2
                        except:
                            pass
                    logging.warning(f"API制限により{wait_time}秒待機中... (試行 {attempt + 1}/{max_retries})")
                    await asyncio.sleep(wait_time)
                    continue
                else:
                    raise HTTPException(
                        status_code=429,
                        detail=f"APIクォータを超過しました。しばらく時間をおいてから再試行してください。"
                    )
            else:
                raise e
    raise HTTPException(status_code=500, detail="最大リトライ回数を超えました。")

# ★ (修正) 低品質アップロード用の SimpleDocumentProcessor を復活
class SimpleDocumentProcessor:
    """
    メモリを消費しない、単純なテキスト抽出とチャンキングを行うクラス。
    unstructured や 親子チャンキング は使用しない。
    """
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        # ユーザーが指定した 1000/200 で分割するスプリッター
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "、", " "]
        )
        logging.info(f"SimpleDocumentProcessor (Chunk: {chunk_size}/{chunk_overlap}) が初期化されました。")

    def _extract_text(self, filename: str, content: bytes) -> str:
        """ファイルタイプに応じてテキストを抽出する"""
        text = ""
        try:
            if filename.endswith(".pdf"):
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    text += page.extract_text()
                logging.info(f".pdf から {len(text)} 文字を抽出 (PyPDF2)")
            
            elif filename.endswith(".docx"):
                doc = DocxDocument(io.BytesIO(content))
                for para in doc.paragraphs:
                    text += para.text + "\n"
                logging.info(f".docx から {len(text)} 文字を抽出 (python-docx)")
            
            elif filename.endswith(".txt"):
                text = content.decode('utf-8')
                logging.info(f".txt から {len(text)} 文字を抽出")
            
            else:
                logging.warning(f"未対応のファイル形式: {filename}")
                
            return re.sub(r'\s+', ' ', text).strip() # 空白を正規化
        
        except Exception as e:
            logging.error(f"テキスト抽出エラー ({filename}): {e}")
            return ""

    def process_and_chunk(self, filename: str, content: bytes, category: str, collection_name: str) -> List[LangChainDocument]:
        """
        1. テキストを抽出
        2. 1000/200 でチャンキング
        3. メタデータを付与 (★親コンテキストは持たない)
        """
        # 1. テキスト抽出
        full_text = self._extract_text(filename, content)
        if not full_text:
            return []

        # 2. チャンキング
        chunks = self.splitter.split_text(full_text)
        
        # 3. メタデータ付与
        docs = []
        for chunk_text in chunks:
            # ★★★ 親子チャンキングではないため、'parent_content' は持たない ★★★
            metadata = {
                "source": filename,
                "collection_name": collection_name,
                "category": category,
                "element_type": "SimpleChunk", # 単純なチャンク
                # "parent_content" はここには無い
            }
            docs.append(LangChainDocument(page_content=chunk_text, metadata=metadata))
            
        logging.info(f"{filename} から {len(docs)} 件の単純チャンクを生成しました。")
        return docs

# (DocumentProcessor クラス 削除済み)

# (WebScraper クラス 削除済み)

class FeedbackManager:
    """フィードバック管理クラス（簡素化版）"""
    def __init__(self):
        self.feedback_file = os.path.join(BASE_DIR, "feedback.json")

    def save_feedback(self, feedback_id: str, rating: str, comment: str = ""):
        """フィードバックを保存"""
        try:
            feedback_data = []
            if os.path.exists(self.feedback_file):
                with open(self.feedback_file, 'r', encoding='utf-8') as f:
                    feedback_data = json.load(f)
            feedback_entry = {
                "id": feedback_id,
                "rating": rating,
                "comment": comment,
                "timestamp": datetime.now(JST).isoformat()
            }
            feedback_data.append(feedback_entry)
            with open(self.feedback_file, 'w', encoding='utf-8') as f:
                json.dump(feedback_data, f, ensure_ascii=False, indent=2)
            logging.info(f"フィードバック保存完了: {feedback_id} - {rating}")
        except Exception as e:
            logging.error(f"フィードバック保存エラー: {e}")
            raise

    def get_feedback_stats(self) -> Dict[str, Any]:
        """フィードバック統計を取得"""
        try:
            if not os.path.exists(self.feedback_file):
                return {"total": 0, "resolved": 0, "not_resolved": 0, "rate": 0}
            with open(self.feedback_file, 'r', encoding='utf-8') as f:
                feedback_data = json.load(f)
            total = len(feedback_data)
            resolved = sum(1 for fb in feedback_data if fb['rating'] == 'resolved')
            not_resolved = total - resolved
            rate = (resolved / total * 100) if total > 0 else 0
            return {
                "total": total,
                "resolved": resolved,
                "not_resolved": not_resolved,
                "rate": round(rate, 1)
            }
        except Exception as e:
            logging.error(f"フィードバック統計取得エラー: {e}")
            return {"total": 0, "resolved": 0, "not_resolved": 0, "rate": 0}

# 8. lifespan関数を更新
@asynccontextmanager
async def lifespan(app: FastAPI):
    """認証システムを含むアプリケーションのライフサイクル管理"""
    # 'g_category_fallbacks' を global 宣言から削除
    # ★ (修正) simple_processor を global に追加
    global db_client, settings_manager, simple_processor
    logging.info("--- アプリケーション起動処理開始(認証システム対応) ---")

    settings_manager = SettingsManager()
    
    # ★ (修正) simple_processor を初期化
    simple_processor = SimpleDocumentProcessor(chunk_size=1000, chunk_overlap=200)

    if SUPABASE_URL and SUPABASE_KEY:
        try:
            db_client = SupabaseClientManager(url=SUPABASE_URL, key=SUPABASE_KEY)
            logging.info("Supabaseクライアントの初期化完了。")



        except Exception as e:
            logging.error(f"Supabase初期化エラー: {e}")
    else:
        logging.warning("Supabaseの環境変数が設定されていません。")

    yield

    logging.info("--- アプリケーション終了処理(認証システム対応) ---")

app = FastAPI(lifespan=lifespan)
app.add_middleware(SessionMiddleware, secret_key=APP_SECRET_KEY)
from prometheus_fastapi_instrumentator import Instrumentator
Instrumentator().instrument(app).expose(app)

# --- グローバルインスタンス ---
feedback_manager = FeedbackManager()
# ★ (修正) simple_processor のインスタンス化は lifespan に移動
# (document_processor 削除済み)
# (web_scraper 削除済み)
# (scrape_parent_splitter, scrape_child_splitter 削除済み)

# --- データモデル定義 ---
class ChatQuery(BaseModel):
    query: str
    model: str = "gemini-2.5-flash"
    embedding_model: str = "text-embedding-004"
    top_k: int = 5
    collection: str = ACTIVE_COLLECTION_NAME

class ClientChatQuery(BaseModel):
    query: str

# (ScrapeRequest 削除済み)

# (重複したChatQueryを修正 -> FeedbackRequest)
class FeedbackRequest(BaseModel):
    feedback_id: str
    rating: str
    comment: str = ""

class AIResponseFeedbackRequest(BaseModel):
    user_question: str
    ai_response: str
    rating: str  # 'good' or 'bad'

class Settings(BaseModel):
    model: Optional[str] = None
    collection: Optional[str] = None
    embedding_model: Optional[str] = None
    top_k: Optional[int] = None

# WebSocket接続管理

manager = ConnectionManager()

# --------------------------------------------------------------------------
# 5. APIエンドポイント定義
# --------------------------------------------------------------------------
@app.post("/feedback/ai-response")
async def save_ai_response_feedback(feedback: AIResponseFeedbackRequest):
    try:
        # feedback.comment にAI回答 or ユーザーコメントを格納
        # created_atはPython側でISOフォーマットのタイムスタンプを格納
        # ratingは good/bad など
        db_client.client.table("anonymous_comments").insert({
            "comment": f"Q: {feedback.user_question}\nA: {feedback.ai_response}",  # 質問・回答をまとめてテキスト化
            "created_at": datetime.now(JST).isoformat(),
            "rating": feedback.rating
        }).execute()
        return {"message": "AI回答へのフィードバックを保存しました"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# main.py のどこか（APIエンドポイントが定義されているエリア）に追加

@app.get("/config")
def get_config():
    """
    フロントエンドが必要とする公開可能な設定（環境変数）を返す
    """
    return {
        "supabase_url": os.getenv("SUPABASE_URL"),
        "supabase_anon_key": os.getenv("SUPABASE_ANON_KEY")
    }

# --- 認証とHTML提供 (Auth0用) ---
# main.pyの修正箇所（ルート定義部分のみ）
# 既存の重複したルート定義を削除し、以下に置き換えてください

# --- 認証とHTML提供 (Auth0用) ---
@app.get('/login')
async def login_auth0(request: Request):
    if 'auth0' not in oauth._clients:
        raise HTTPException(status_code=500, detail="Auth0 is not configured.")
    return await oauth.auth0.authorize_redirect(request, request.url_for('auth'))

@app.get('/auth')
async def auth(request: Request):
    """
    Auth0からのコールバックを処理し、ユーザー情報に基づいて
    適切なページ（/admin または /）にリダイレクトする。
    """
    if 'auth0' not in oauth._clients:
        raise HTTPException(status_code=500, detail="Auth0 is not configured.")
    
    try:
        token = await oauth.auth0.authorize_access_token(request)
    except Exception as e:
        logging.error(f"Auth0 access token error: {e}")
        return RedirectResponse(url='/login')

    if userinfo := token.get('userinfo'):
        # ユーザー情報をセッションに保存
        request.session['user'] = dict(userinfo)
        user_email = userinfo.get('email', '').lower() # 小文字に変換
        
        # 許可リストを小文字に
        super_admin_emails_lower = [email.lower() for email in SUPER_ADMIN_EMAILS]
        allowed_emails_lower = [email.lower() for email in ALLOWED_CLIENT_EMAILS]

        # --- 権限に基づくリダイレクト判定 ---
        
        # 1. まず、管理者(SUPER_ADMIN)か？ (最優先)
        #    管理者は /admin にリダイレクト
        if user_email in super_admin_emails_lower:
            return RedirectResponse(url='/admin')

        # 2. 次に、クライアント(学生など)か？
        #    クライアントは / (学生用画面) にリダイレクト
        elif user_email in allowed_emails_lower:
            return RedirectResponse(url='/')
        
        # 3. 上記のいずれにも該当しない場合は、許可されていないユーザー
        #    ログアウトさせてアクセスを拒否
        else:
            logging.warning(f"Unauthorized login attempt by: {user_email}")
            return RedirectResponse(url='/logout')
            
    # Auth0からuserinfoが取得できなかった場合
    logging.error("Failed to get userinfo from Auth0.")
    return RedirectResponse(url='/login')

@app.get('/logout')
async def logout(request: Request):
    request.session.pop('user', None)
    if not all([AUTH0_DOMAIN, AUTH0_CLIENT_ID]):
        return RedirectResponse(url='/')
    return RedirectResponse(f"https://{AUTH0_DOMAIN}/v2/logout?returnTo={request.url_for('serve_client')}&client_id={AUTH0_CLIENT_ID}")

# クライアント用ページ（学生用）
@app.get("/", response_class=FileResponse)
async def serve_client(request: Request, user: dict = Depends(require_auth_client)):
    return FileResponse(os.path.join(BASE_DIR, "client.html"))

# 管理者用メインページ
@app.get("/admin", response_class=FileResponse)
async def serve_admin(request: Request, user: dict = Depends(require_auth)):
    return FileResponse(os.path.join(BASE_DIR, "admin.html"))

# ★★★ 重要: DB管理ページ（管理者のみアクセス可能）★★★
# 重複を削除し、この1つだけ残します
@app.get("/DB.html", response_class=FileResponse)
async def serve_db_page(request: Request, user: dict = Depends(require_auth)):
    """DB管理ページ(DB.html)を提供するためのエンドポイント"""
    db_path = os.path.join(BASE_DIR, "DB.html")
    if not os.path.exists(db_path):
        raise HTTPException(status_code=404, detail="DB.html not found")
    return FileResponse(db_path)

# CSS提供
@app.get("/style.css", response_class=FileResponse)
async def serve_css():
    """統一されたCSSファイル(style.css)を提供するためのエンドポイント"""
    return FileResponse(os.path.join(BASE_DIR, "style.css"))

# フィードバック統計ページ
@app.get("/feedback-stats", response_class=FileResponse)
async def serve_feedback_stats(request: Request, user: dict = Depends(require_auth)):
    return FileResponse(os.path.join(BASE_DIR, "feedback_stats.html"))

# Faviconエラー回避
@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    return Response(status_code=204)

# --- ステータス確認 ---
@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "database": db_client.get_db_type() if db_client else "uninitialized"
    }

@app.get("/gemini/status")
async def gemini_status():
    if not GEMINI_API_KEY:
        return {"connected": False, "detail": "API key not configured"}
    try:
        models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        return {"connected": True, "models": models}
    except Exception as e:
        return {"connected": False, "detail": str(e)}

@app.api_route("/healthz", methods=["GET", "HEAD"])
def health_check_k8s():
    return {"status": "ok"}

# --- コレクション管理API ---
@app.get("/collections")
async def get_collections():
    return [{"name": ACTIVE_COLLECTION_NAME, "count": db_client.count_chunks_in_collection(ACTIVE_COLLECTION_NAME) if db_client else 0}]

@app.post("/collections")
async def create_collection(request: dict):
    return {"message": f"コレクション「{ACTIVE_COLLECTION_NAME}」は既に存在しています"}

@app.delete("/collections/{collection_name}")
async def delete_collection(collection_name: str):
    if collection_name == ACTIVE_COLLECTION_NAME:
        raise HTTPException(status_code=400, detail="このコレクションは削除できません")
    return {"message": "コレクションが見つかりません"}

# --- ナレッジ管理API ---
# --- ナレッジ管理API ---
@app.get("/collections/{collection_name}/documents")
async def get_documents(collection_name: str):
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    return {
        "documents": db_client.get_documents_by_collection(collection_name),
        "count": db_client.count_chunks_in_collection(collection_name)
    }

# --- ドキュメント管理API（新規追加） ---
@app.get("/api/documents/all")
async def get_all_documents(
    user: dict = Depends(require_auth),
    page: int = Query(1, ge=1),
    limit: int = Query(100, ge=1, le=1000),
    search: Optional[str] = Query(None),
    category: Optional[str] = Query(None)
):
    """
    全ドキュメントをページネーション、検索、カテゴリフィルタ対応で取得（管理者のみ）
    """
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    try:
        # --- 1. ベースとなるクエリを構築 ---
        # (SupabaseのPostgRESTビルダーを使って、動的にクエリを組み立てます)
        query = db_client.client.table("documents")
        count_query = db_client.client.table("documents")
        # --- 2. フィルタ条件の適用 ---
        if category:
            query = query.eq("metadata->>category", category)
            count_query = count_query.eq("metadata->>category", category)

        if search:
                # 1. 検索語自体に含まれる " を "" にエスケープ
                safe_search = search.replace('"', '""')
                
                # 2. "*" ワイルドカードを付与し、パターン全体を二重引用符 " で囲む
                search_term = f"*{safe_search}*"
                
                # 3. .or_() に渡す単一の文字列を構築
                # ★★★ 修正点 ★★★
                # metadata->>category の検索が500エラーを引き起こすため、
                # 検索対象を content と source のみ に絞ります。
                or_filter_string = (
                    f"content.ilike.{search_term},"
                    # f"metadata->>source.ilike.{search_term}"
                )
                
                # 4. クエリに適用
                query = query.or_(or_filter_string)
                count_query = count_query.or_(or_filter_string)

        # --- 3. 総件数の取得 (フィルタ適用後) ---
        # .select("id", count='exact') で、実際のデータを取得せず件数だけを取得
        count_response = count_query.select("id", count='exact').execute()
        total_records = count_response.count or 0

        # --- 4. ページ指定されたデータの取得 ---
        # Supabase (PostgREST) の range は limit/offset 方式
        offset = (page - 1) * limit
        # .range(from, to) ※ to は inclusive
        data_response = query.select("*").order("id", desc=True).range(offset, offset + limit - 1).execute()

        return {
            "documents": data_response.data or [],
            "total": total_records,
            "page": page,
            "limit": limit
        }

    except Exception as e:
        logging.error(f"ドキュメント一覧(ページネーション)取得エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/documents/{doc_id}")
async def get_document_by_id(doc_id: int, user: dict = Depends(require_auth)):
    """特定のドキュメントを取得（管理者のみ）"""
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    try:
        result = db_client.client.table("documents").select("*").eq("id", doc_id).execute()
        if not result.data:
            raise HTTPException(status_code=404, detail="ドキュメントが見つかりません")
        return result.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"ドキュメント取得エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/documents/{doc_id}")
async def update_document(doc_id: int, request: Dict[str, Any], user: dict = Depends(require_auth)):
    """ドキュメントを更新（管理者のみ）。content更新時にベクトルも再生成する。"""
    if not db_client or not settings_manager:
        raise HTTPException(503, "DBまたは設定マネージャーが初期化されていません")
    try:
        update_data = {}

        # 1. contentが更新される場合は、ベクトルも再生成する
        if "content" in request:
            new_content = request["content"]
            update_data["content"] = new_content
            
            # 設定マネージャーから現在のエンベディングモデルを取得
            embedding_model = settings_manager.settings.get("embedding_model", "text-embedding-004")
            
            logging.info(f"ドキュメント {doc_id} のコンテンツが変更されたため、ベクトルを再生成します...")
            try:
                # 新しいコンテンツでベクトルを生成
                embedding_response = genai.embed_content(
                    model=embedding_model,
                    content=new_content
                )
                update_data["embedding"] = embedding_response["embedding"]
                logging.info(f"ドキュメント {doc_id} のベクトル再生成が完了しました。")
            except Exception as e:
                # APIのレート制限などを考慮
                if "429" in str(e) or "quota" in str(e).lower():
                    logging.warning("ベクトル再生成でAPI制限に達しました。30秒待機します。")
                    await asyncio.sleep(30)
                    embedding_response = genai.embed_content(
                        model=embedding_model,
                        content=new_content
                    )
                    update_data["embedding"] = embedding_response["embedding"]
                else:
                    logging.error(f"ベクトル再生成エラー: {e}")
                    raise HTTPException(status_code=500, detail=f"ベクトル再生成中にエラーが発生しました: {e}")

        # 2. メタデータは独立して更新可能
        if "metadata" in request:
            update_data["metadata"] = request["metadata"]

        if not update_data:
            raise HTTPException(status_code=400, detail="更新するデータがありません")
        
        result = db_client.client.table("documents").update(update_data).eq("id", doc_id).execute()
        
        if not result.data:
            raise HTTPException(status_code=404, detail="ドキュメントが見つかりません")
        
        logging.info(f"ドキュメント {doc_id} を更新しました（管理者: {user.get('email')}）")
        return {"message": "ドキュメントを更新しました", "document": result.data[0]}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"ドキュメント更新エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: int, user: dict = Depends(require_auth)):
    """ドキュメントを削除（管理者のみ）"""
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    try:
        result = db_client.client.table("documents").delete().eq("id", doc_id).execute()
        
        if not result.data:
            raise HTTPException(status_code=404, detail="ドキュメントが見つかりません")
        
        logging.info(f"ドキュメント {doc_id} を削除しました（管理者: {user.get('email')}）")
        return {"message": "ドキュメントを削除しました"}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"ドキュメント削除エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# ★ (修正) 低品質アップロード用の /upload エンドポイントを復活
@app.post("/upload")
async def upload_document(
    file: UploadFile = File(...), 
    category: str = Form("その他"), 
    user: dict = Depends(require_auth)
):
    """
    ファイルを受け取り、(単純な)テキスト抽出とチャンキングを行い、
    ベクトル化してDBに挿入する (低メモリ版)
    """
    if not db_client or not settings_manager or not simple_processor:
        raise HTTPException(503, "システムが初期化されていません (DB, Settings, or Processor)")

    try:
        filename = file.filename
        content = await file.read()
        
        logging.info(f"ファイルアップロード受信: {filename} (カテゴリ: {category})")

        # 1. (新しい) SimpleDocumentProcessor で処理
        collection_name = settings_manager.settings.get("collection", ACTIVE_COLLECTION_NAME)
        docs_to_embed = simple_processor.process_and_chunk(filename, content, category, collection_name)
        
        if not docs_to_embed:
            raise HTTPException(status_code=400, detail="ファイルからテキストを抽出できませんでした。")

        embedding_model = settings_manager.settings.get("embedding_model", "text-embedding-004")
        total_chunks = len(docs_to_embed)
        logging.info(f"{total_chunks} 件のチャンクをベクトル化・挿入します...")

        # 2. ループしてベクトル化 & DB挿入
        count = 0
        for doc in docs_to_embed:
            try:
                # 検索対象のテキスト (page_content) をベクトル化
                embedding_response = genai.embed_content(
                    model=embedding_model, 
                    content=doc.page_content
                )
                embedding = embedding_response["embedding"]
                
                # DBにはチャンクテキスト、ベクトル、メタデータ(親なし)を格納
                db_client.insert_document(
                    content=doc.page_content, 
                    embedding=embedding, 
                    metadata=doc.metadata
                )
                count += 1
                
                await asyncio.sleep(1) # APIレート制限対策
            
            except Exception as e:
                if "429" in str(e) or "quota" in str(e).lower():
                    logging.warning("埋め込み生成でAPI制限に達しました。30秒待機します。")
                    await asyncio.sleep(30)
                    # (再試行)
                    embedding_response = genai.embed_content(model=embedding_model, content=doc.page_content)
                    embedding = embedding_response["embedding"]
                    db_client.insert_document(doc.page_content, embedding, doc.metadata)
                else:
                    logging.error(f"チャンク処理エラー ({filename}): {e}")
                    continue # 次のチャンクへ

        logging.info(f"ファイル処理完了: {filename} ({count}/{total_chunks}件のチャンクをDBに挿入)")
        return {"chunks": count, "filename": filename, "total": total_chunks}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"ファイルアップロード処理エラー: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

# (@app.post("/scrape") 削除済み)
    
 # --- チャットAPI（RAG対応） ---

# 686行目から734行目までの不完全なenhanced_chat_logic関数を削除

# その後、738行目付近に以下の関数を追加(グローバルスコープ)
def get_or_create_session_id(request: Request) -> str:
    """セッションIDを取得または新規作成"""
    session_id = request.session.get('chat_session_id')
    if not session_id:
        session_id = str(uuid.uuid4())
        request.session['chat_session_id'] = session_id
    return session_id

def add_to_history(session_id: str, role: str, content: str):
    """チャット履歴に追加"""
    # ---------- ここを一時的に無効化しています ----------
    # 会話の記録を完全に停止したい（コメントアウトで一時無効化）
    # 元に戻す際は、下のコードをコメント解除してください。
    #
    # chat_histories[session_id].append({
    #     "role": role,
    #     "content": content
    # })
    # if len(chat_histories[session_id]) > MAX_HISTORY_LENGTH:
    #     chat_histories[session_id] = chat_histories[session_id][-MAX_HISTORY_LENGTH:]
    return
# ------------------------------------------------------

def get_history(session_id: str) -> List[Dict[str, str]]:
    """チャット履歴を取得"""
    return chat_histories.get(session_id, [])


def clear_history(session_id: str):
    """チャット履歴をクリア"""
    if session_id in chat_histories:
        del chat_histories[session_id]


# ===============================
# 正しい enhanced_chat_logic 定義
# ===============================
async def enhanced_chat_logic(request: Request, chat_req: ChatQuery):
    """RAG + フォールバック対応のチャット処理"""
    user_input = chat_req.query.strip()
    feedback_id = str(uuid.uuid4())
    session_id = get_or_create_session_id(request)

    yield f"data: {json.dumps({'feedback_id': feedback_id})}\n\n"

    try:
        if not all([db_client, GEMINI_API_KEY]):
            yield f"data: {json.dumps({'content': 'システムが利用できません。管理者にお問い合わせください。'})}\n\n"
            return

        # =======================
        # ベクトル検索処理
        # =======================
        STRICT_THRESHOLD = 0.80
        RELATED_THRESHOLD = 0.75
        search_results = []
        relevant_docs = []

        try:
            query_embedding_response = genai.embed_content(
                model=chat_req.embedding_model,
                content=user_input
            )
            query_embedding = query_embedding_response["embedding"]

            if db_client:
                search_results = db_client.search_documents_by_vector(
                    collection_name=chat_req.collection,
                    embedding=query_embedding,
                    match_count=chat_req.top_k
                )

            logging.info(f"検索結果件数: {len(search_results)}")

        except Exception as e:
            logging.error(f"ベクトル検索エラー: {e}")
            search_results = []

        # 類似度フィルタリング
        strict_docs = [d for d in search_results if d.get('similarity', 0) >= STRICT_THRESHOLD]
        related_docs = [d for d in search_results if RELATED_THRESHOLD <= d.get('similarity', 0) < STRICT_THRESHOLD]
        relevant_docs = strict_docs + related_docs
        # --- ▼ ログ出力の追加 ▼ ---
        if relevant_docs:
           logging.info(f"--- Stage 1 RAG ヒット (上位 {len(relevant_docs)}件) ---")
        for doc in relevant_docs:
                # ログに出力したい情報を doc (辞書) から取得
                doc_id = doc.get('id', 'N/A')
                doc_source = doc.get('metadata', {}).get('source', 'N/A')
                doc_similarity = doc.get('similarity', 0)
                
                # content の内容をプレビューとして取得 (冒頭50文字 + 改行をスペースに置換)
                doc_content_preview = doc.get('content', '')[:50].replace('\n', ' ') + "..."
                
                # ログ出力を強化
                logging.info(f"  [ID: {doc_id}] [Sim: {doc_similarity:.4f}] (Source: {doc_source}) Content: '{doc_content_preview}'")
        # --- ▲ ログ出力の追加 ▲ ---

        # =======================
        # コンテキスト生成と回答生成
        # =======================
        if relevant_docs:
            # context を構築するための空リストを用意
            context_parts = []
            
            for d in relevant_docs:
                # 1. 元の source (ファイル名) を取得
                source_name = d.get('metadata', {}).get('source', '不明')
                
                # 2. ★★★ ここで表示名をマッピング ★★★
                if source_name == 'output_gakubu.txt':
                    display_source = '履修要項2024'
                # (必要なら、他のファイル名もマッピングできます)
                # elif source_name == 'another_file.pdf':
                #     display_source = '学生生活の手引き'
                else:
                    display_source = source_name # マッピング対象外はそのまま
                
                # 3. 構築
                context_parts.append(
                    f"<document source='{display_source}'>{d.get('content', '')}</document>"
                )
            
            # 最後に context を結合
            context = "\n\n".join(context_parts)

            # --- RAGプロンプトとAI呼び出し (if ブロックの内部) ---
            prompt = f"""あなたは札幌学院大学の学生サポートAIです。  
以下のルールに従ってユーザーの質問に答えてください。

# ルール
1. 回答は <context> 内の情報（大学公式情報）のみに基づいてください。
2. <context> に質問と「完全に一致する答え」が見つからない場合でも、「関連する可能性のある情報」（例：質問は「大会での欠席」だが、資料には「病欠」について記載がある場合）が見つかった場合は、その情報を回答してください。
3. （ルール#2 に基づき）関連情報で回答した場合は、回答の最後に必ず以下の「注意書き」を加えてください。
   「※これは関連情報であり、ご質問の意図と完全に一致しない可能性があります。詳細は大学の公式窓口にご確認ください。」
4. 出典を引用する場合は、使用した情報の直後に `[出典: ...]` を付けてください。
5. 大学固有の情報を推測して答えてはいけません。
6. **特に重要**: <context> には必ず関連情報が含まれています。その情報を使って回答すること。「見つかりませんでした」と答えてはいけません。

# 出力形式
- 学生に分かりやすい「です・ます調」で回答すること。
- 箇条書きや見出しを活用して整理すること。
- <context> 内にURLがあれば「参考URL:」として末尾にまとめること。

<context>
{context}
</context>

<query>
{user_input}
</query>

---
[あなたの回答]
回答:
"""

            # ★★★ 安全フィルターを無効化する設定 ★★★
            # 大学の履修要項という安全なテキストでも、組み合わせによって
            # フィルターに抵触してしまう誤検知(False Positive)が観測されたため。
            safety_settings = {
                HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
                HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            }
            
            model = genai.GenerativeModel(
                chat_req.model,
                safety_settings=safety_settings # ★★★ この設定を追加 ★★★
            )
            response_text = ""
            try:
                stream = await safe_generate_content(model, prompt, stream=True)
                async for chunk in stream:
                    if chunk.text:
                        response_text += chunk.text
            # ... (876行目付近)
            except Exception as e:
                logging.error(f"生成エラー: {e}")
                response_text = "回答の生成中にエラーが発生しました。"

            full_response = format_urls_as_links(response_text.strip() or "回答を生成できませんでした。")
            
            # ★★★ 修正箇所: 成功したRAGの回答をここで yield する ★★★
            add_to_history(session_id, "user", user_input)
            add_to_history(session_id, "assistant", response_text) # 生のテキストを記録
            yield f"data: {json.dumps({'content': full_response})}\n\n"
            # ★★★ 修正ここまで ★★★

        else:
            # --- フォールバック処理 (Stage 2: Q&Aベクトル検索) ---
            logging.info(f"Stage 1 RAG 失敗。Stage 2 (Q&Aベクトル検索) を実行します。")

            try:
                # Stage 1 で使用した query_embedding を再利用
                fallback_results = db_client.search_fallback_qa(
                    embedding=query_embedding,
                    match_count=1  # 最も近いQ&Aを1つだけ取得
                )

                if fallback_results:
                    best_match = fallback_results[0]
                    FALLBACK_SIMILARITY_THRESHOLD = 0.59  # フォールバック用のしきい値

                    if best_match.get('similarity', 0) >= FALLBACK_SIMILARITY_THRESHOLD:
                        logging.info(
                            f"Stage 2 RAG 成功。類似Q&Aを回答します (Similarity: {best_match['similarity']:.2f})"
                        )
                        fallback_response = f"""データベースに直接の情報は見つかりませんでしたが、関連する「よくあるご質問」がありましたのでご案内します。

---
{best_match['content']}
"""
                        full_response = format_urls_as_links(fallback_response)
                    else:
                        logging.info(
                            f"Stage 2 RAG 失敗。類似するQ&Aが見つかりませんでした (Best Similarity: {best_match.get('similarity', 0):.2f})"
                        )
                        fallback_response = "申し訳ありませんが、ご質問に関連する情報がデータベース（Q&Aを含む）に見つかりませんでした。大学公式サイトをご確認いただくか、学生支援課までお問い合わせください。"
                        full_response = format_urls_as_links(fallback_response)
                else:
                    logging.info("Stage 2 RAG 失敗。Q&Aデータベースが空か、検索エラーです。")
                    fallback_response = "申し訳ありませんが、ご質問に関連する情報が見つかりませんでした。大学公式サイトをご確認いただくか、学生支援課までお問い合わせください。"
                    full_response = format_urls_as_links(fallback_response)

            except Exception as e_fallback:
                logging.error(f"Stage 2 (Q&A検索) でエラーが発生: {e_fallback}")
                fallback_response = "申し訳ありません。現在、関連情報の検索中にエラーが発生しました。時間をおいて再度お試しください。"
                full_response = format_urls_as_links(fallback_response)

            add_to_history(session_id, "user", user_input)
            add_to_history(session_id, "assistant", fallback_response)  # 生のテキストを記録

            yield f"data: {json.dumps({'content': full_response})}\n\n"

        # 最後のフィードバック出力
        yield f"data: {json.dumps({'show_feedback': True, 'feedback_id': feedback_id})}\n\n"

    except Exception as e:
        err = f"エラーが発生しました: {e}"
        logging.error(f"チャットロジック全体エラー: {err}")
        yield f"data: {json.dumps({'content': err})}\n\n"


# 履歴管理用エンドポイント
@app.get("/chat/history")
async def get_chat_history(request: Request, user: dict = Depends(require_auth_client)):
    """現在のセッションのチャット履歴を取得"""
    session_id = get_or_create_session_id(request)
    history = get_history(session_id)
    return {"history": history}

@app.delete("/chat/history")
async def delete_chat_history(request: Request, user: dict = Depends(require_auth_client)):
    """現在のセッションのチャット履歴を削除"""
    session_id = get_or_create_session_id(request)
    clear_history(session_id)
    return {"message": "履歴をクリアしました"}

@app.post("/chat")
async def chat_endpoint(request: Request, query: ChatQuery):
    return StreamingResponse(enhanced_chat_logic(request, query), media_type="text/event-stream")

# 7. 既存のチャットエンドポイントを認証必須に変更
@app.post("/chat_for_client")
async def chat_for_client_auth(request: Request, query: ClientChatQuery, user: dict = Depends(require_auth_client)):
    """認証されたクライアント用チャットエンドポイント"""
    if not settings_manager:
        raise HTTPException(503, "Settings manager not initialized")

    logging.info(f"Chat request from user: {user.get('email', 'N/A')}")

    chat_query = ChatQuery(
        query=query.query,
        model=settings_manager.settings.get("model", "gemini-2.5-flash"),
        embedding_model=settings_manager.settings.get("embedding_model", "text-embedding-004"),
        top_k=settings_manager.settings.get("top_k", 5),
        collection=settings_manager.settings.get("collection", ACTIVE_COLLECTION_NAME)
    )
    return StreamingResponse(enhanced_chat_logic(request, chat_query), media_type="text/event-stream")


# --- フィードバックAPI ---
@app.post("/feedback")
async def save_feedback(feedback: FeedbackRequest):
    try:
        feedback_manager.save_feedback(feedback.feedback_id, feedback.rating, feedback.comment)
        return {"message": "フィードバックを保存しました"}
    except Exception as e:
        logging.error(f"フィードバック保存エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Q&A (category_fallbacks) 管理API ---
# main.py のAPIエンドポイント定義エリア (例: 1150行目あたり) に追加してください

@app.get("/api/fallbacks")
async def get_all_fallbacks(user: dict = Depends(require_auth)):
    """Q&A(フォールバック)をすべて取得"""
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    try:
        result = db_client.client.table("category_fallbacks").select("*").order("id", desc=True).execute()
        return {"fallbacks": result.data or []}
    except Exception as e:
        logging.error(f"Q&A一覧取得エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# [main.py] 1162行目あたり

# [main.py] 1162行目あたり

@app.post("/api/fallbacks")
async def create_fallback(request: Dict[str, Any], user: dict = Depends(require_auth)):
    """新しいQ&Aを作成（保存時に自動でベクトル化）"""
    # 依存関係に settings_manager を追加
    if not db_client or not settings_manager: 
        raise HTTPException(503, "DBまたは設定マネージャーが初期化されていません")
    
    try:
        new_qa_text = request.get("static_response", "")
        category_name = request.get("category_name") 

        if not new_qa_text:
            raise HTTPException(status_code=400, detail="static_response (Q&Aテキスト) は必須です")
        
        if not category_name:
            raise HTTPException(status_code=400, detail="category_name は必須です")

        # --- ここから自動ベクトル化ロジック ---
        embedding = None
        try:
            # 設定マネージャーからエンベディングモデルを取得
            embedding_model = settings_manager.settings.get("embedding_model", "text-embedding-004")
            logging.info(f"新規Q&Aのベクトルを生成します...")
            
            # ベクトルを生成
            embedding_response = genai.embed_content(
                model=embedding_model,
                content=new_qa_text
            )
            embedding = embedding_response["embedding"]
            logging.info(f"新規Q&Aのベクトル生成が完了しました。")
        
        except Exception as e:
            # レート制限などで失敗しても、テキストの保存は続行する (embedding = None のまま)
            logging.error(f"新規Q&Aのベクトル生成エラー: {e}")
            logging.warning(f"ベクトル化に失敗しましたが、テキストは保存します。")
        # --- ここまで ---

        # embedding を挿入データに含める
        insert_data = {
            "static_response": new_qa_text,
            "category_name": category_name,
            "url_to_summarize": request.get("url_to_summarize"), # (古いカラムが残っている場合)
            "embedding": embedding  # (None またはベクトル)
        }
        
        result = db_client.client.table("category_fallbacks").insert(insert_data).execute()
        
        if not result.data:
            raise HTTPException(status_code=500, detail="Q&Aの作成に失敗しました")

        logging.info(f"新規Q&A {result.data[0]['id']} を作成しました（管理者: {user.get('email')}）")
        
        # メッセージを変更
        message = "新しいQ&Aを保存し、ベクトル化も完了しました。" if embedding else "新しいQ&Aを保存しました（ベクトル化には失敗）"
        return {"message": message, "fallback": result.data[0]}
    
    except HTTPException:
        raise # 400エラーなどをそのまま返す
    except Exception as e:
        # DB制約エラーのハンドリング
        if "23502" in str(e) and "category_name" in str(e):
             raise HTTPException(status_code=400, detail="category_name は必須です (DB Error 23502)")
        logging.error(f"Q&A作成エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# [main.py] 1190行目あたり

@app.put("/api/fallbacks/{qa_id}")
async def update_fallback(qa_id: int, request: Dict[str, Any], user: dict = Depends(require_auth)):
    """Q&Aを更新（テキスト変更時に自動でベクトル化）"""
    if not db_client or not settings_manager:
        raise HTTPException(503, "DBまたは設定マネージャーが初期化されていません")
    try:
        update_data = {}
        
        # 1. static_response (Q&Aテキスト) が更新されるかチェック
        if "static_response" in request:
            new_content = request["static_response"]
            update_data["static_response"] = new_content
            
            # --- 自動ベクトル化 ---
            embedding_model = settings_manager.settings.get("embedding_model", "text-embedding-004")
            logging.info(f"Q&A {qa_id} のテキストが変更されたため、ベクトルを再生成します...")
            try:
                embedding_response = genai.embed_content(
                    model=embedding_model,
                    content=new_content
                )
                update_data["embedding"] = embedding_response["embedding"]
                logging.info(f"Q&A {qa_id} のベクトル再生成が完了しました。")
            except Exception as e:
                logging.error(f"Q&Aベクトル再生成エラー: {e}")
                update_data["embedding"] = None
                logging.warning(f"Q&A {qa_id} のベクトル化に失敗しましたが、テキストは更新します。")

        # 2. url_to_summarize も更新可能 (古いカラムが残っている場合)
        if "url_to_summarize" in request:
            update_data["url_to_summarize"] = request.get("url_to_summarize")
            
        # ★★★ 3. category_name の更新に対応 ★★★
        if "category_name" in request:
            new_category = request.get("category_name")
            if not new_category or not new_category.strip():
                 raise HTTPException(status_code=400, detail="category_name を空にすることはできません")
            update_data["category_name"] = new_category

        if not update_data:
            raise HTTPException(status_code=400, detail="更新するデータがありません")
        
        result = db_client.client.table("category_fallbacks").update(update_data).eq("id", qa_id).execute()
        
        if not result.data:
            raise HTTPException(status_code=404, detail="Q&Aが見つかりません")
        
        logging.info(f"Q&A {qa_id} を更新しました（管理者: {user.get('email')}）")
        return {"message": "Q&Aを更新しました", "fallback": result.data[0]}
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Q&A更新エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/fallbacks/{qa_id}")
async def delete_fallback(qa_id: int, user: dict = Depends(require_auth)):
    """Q&Aを削除"""
    if not db_client:
        raise HTTPException(503, "DB not initialized")
    try:
        result = db_client.client.table("category_fallbacks").delete().eq("id", qa_id).execute()
        if not result.data:
            raise HTTPException(status_code=404, detail="Q&Aが見つかりません")
        
        logging.info(f"Q&A {qa_id} を削除しました（管理者: {user.get('email')}）")
        return {"message": "Q&Aを削除しました"}
    except Exception as e:
        logging.error(f"Q&A削除エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/fallbacks/vectorize-all")
async def vectorize_all_missing_fallbacks(user: dict = Depends(require_auth)):
    """embedding が NULL のQ&Aをすべてベクトル化する"""
    if not db_client or not settings_manager:
        raise HTTPException(503, "DBまたは設定マネージャーが初期化されていません")

    logging.info(f"全Q&Aのベクトル化処理を開始...（管理者: {user.get('email')}）")
    
    try:
        # 1. embedding が NULL の Q&A をすべて取得
        response = db_client.client.table("category_fallbacks").select("id, static_response").is_("embedding", "null").execute()
        
        if not response.data:
            return {"message": "ベクトル化が必要なQ&Aはありませんでした。"}

        embedding_model = settings_manager.settings.get("embedding_model", "text-embedding-004")
        count = 0
        
        for item in response.data:
            item_id = item['id']
            text_to_vectorize = item['static_response']

            if not text_to_vectorize or not text_to_vectorize.strip():
                logging.warning(f"Q&A ID {item_id}: テキストが空のためスキップします。")
                continue

            try:
                # 2. ベクトル化
                embedding_response = genai.embed_content(
                    model=embedding_model,
                    content=text_to_vectorize
                )
                new_embedding = embedding_response["embedding"]

                # 3. DBを更新
                db_client.client.table("category_fallbacks").update({
                    "embedding": new_embedding
                }).eq("id", item_id).execute()
                
                logging.info(f"Q&A ID {item_id}: ベクトル化完了。")
                count += 1
                await asyncio.sleep(1) # APIレート制限対策

            except Exception as e:
                if "429" in str(e) or "quota" in str(e).lower():
                    logging.warning(f"APIレート制限のため30秒待機します... (ID {item_id})")
                    await asyncio.sleep(30)
                    # 1回だけ再試行
                    embedding_response = genai.embed_content(model=embedding_model, content=text_to_vectorize)
                    new_embedding = embedding_response["embedding"]
                    db_client.client.table("category_fallbacks").update({"embedding": new_embedding}).eq("id", item_id).execute()
                    logging.info(f"Q&A ID {item_id}: (再試行) ベクトル化完了。")
                    count += 1
                else:
                    logging.error(f"Q&A ID {item_id} のベクトル化エラー: {e}")

        logging.info(f"全Q&Aベクトル化処理完了。 {count}件を処理しました。")
        return {"message": f"ベクトル化処理が完了しました。{count}件のQ&Aを更新しました。"}

    except Exception as e:
        logging.error(f"全Q&Aベクトル化処理中にエラーが発生: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/feedback/stats")
async def get_feedback_stats():
    try:
        stats = feedback_manager.get_feedback_stats()
        return stats
    except Exception as e:
        logging.error(f"フィードバック統計取得エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- 設定管理API ---
@app.post("/settings")
async def update_settings(settings: Settings):
    if not settings_manager:
        raise HTTPException(503, "設定マネージャーが初期化されていません")
    try:
        await settings_manager.update_settings(settings.dict(exclude_none=True))
        return {"message": "設定を更新しました"}
    except Exception as e:
        logging.error(f"設定更新エラー: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- WebSocketエンドポイント ---
@app.websocket("/ws/settings")
async def websocket_endpoint(websocket: WebSocket):
    if not settings_manager:
        await websocket.close(code=1011, reason="Settings manager not initialized")
        return
    await settings_manager.add_websocket(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        settings_manager.remove_websocket(websocket)

# --------------------------------------------------------------------------
# 6. 開発用サーバー起動
# --------------------------------------------------------------------------
if __name__ == "__main__":
    port = int(os.getenv("PORT", 10000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)